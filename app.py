import streamlit as st
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from io import BytesIO
import datetime

st.set_page_config(page_title="Gerador de Propostas", layout="wide")

# Logo fixa
st.image("LOGO DGCE.png", width=200)
st.title("📄 Gerador de Propostas Comerciais Técnicas")

# ==========================================================
# FUNÇÕES AUXILIARES
# ==========================================================

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def substituir_placeholders(doc, dados, tabela_itens):
    for p in doc.paragraphs:
        # Logo
        if "{{LOGO}}" in p.text:
            p.text = ""
            run = p.add_run()
            run.add_picture("LOGO DGCE.png", width=Inches(2))

        # Tabela dinâmica no lugar certo
        elif "{{TABELA}}" in p.text:
            p.text = ""  # remove o placeholder

            if tabela_itens:
                # cria elemento de tabela XML
                tbl = OxmlElement('w:tbl')

                # cabeçalho
                tr = OxmlElement('w:tr')
                for h in ["Item", "Incluído", "Não Incluído"]:
                    tc = OxmlElement('w:tc')
                    p_tc = OxmlElement('w:p')
                    r = OxmlElement('w:r')
                    t = OxmlElement('w:t')
                    t.text = h
                    r.append(t)
                    p_tc.append(r)
                    tc.append(p_tc)
                    tr.append(tc)
                tbl.append(tr)

                # linhas da tabela
                for item in tabela_itens:
                    tr = OxmlElement('w:tr')
                    for val in [item["Item"], item["Incluso"], item["Nao_Incluso"]]:
                        tc = OxmlElement('w:tc')
                        p_tc = OxmlElement('w:p')
                        r = OxmlElement('w:r')
                        t = OxmlElement('w:t')
                        t.text = val
                        r.append(t)
                        p_tc.append(r)
                        tc.append(p_tc)
                        tr.append(tc)
                    tbl.append(tr)

                # insere a tabela logo após o parágrafo
                p._element.addnext(tbl)

        # Listas e textos simples
        else:
            for chave, valor in dados.items():
                if f"{{{{{chave}}}}}" in p.text:
                    if chave in ["BENEFICIOS","ESCOPO","OBSERVACOES",
                                 "RESPONSABILIDADES_CONTRATADA","RESPONSABILIDADES_CONTRATANTE"]:
                        itens = [i.strip() for i in valor.split(";") if i.strip()]
                        p.text = p.text.replace(f"{{{{{chave}}}}}", "\n".join(itens))
                    else:
                        p.text = p.text.replace(f"{{{{{chave}}}}}", valor)
    return doc




def gerar_docx(dados, tabela_itens, template_file):
    if template_file is not None:
        doc = Document(template_file)
    else:
        doc = Document("PROJETOS.docx")  # usa o modelo padrão
    doc = substituir_placeholders(doc, dados, tabela_itens)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ==========================================================
# SIDEBAR - ENTRADAS
# ==========================================================
st.sidebar.header("📝 Dados da Proposta")

# Campos obrigatórios
nome_cliente = st.sidebar.text_input("Nome do Cliente")
titulo_projeto = st.sidebar.text_input("Título do Projeto")
valor_total = st.sidebar.text_input("Valor Total (R$)")
prazo_entrega = st.sidebar.text_input("Prazo de Entrega")
escopo_tecnico = st.sidebar.text_area("Escopo Técnico")

# Campos opcionais
ano = st.sidebar.text_input("Ano")
objetivo = st.sidebar.text_area("Objetivo")
beneficios = st.sidebar.text_area("Benefícios (separe por ;)")
referencias = st.sidebar.text_area("Referências")
escopo = st.sidebar.text_area("Escopo (separe por ;)")
observacoes = st.sidebar.text_area("Observações (separe por ;)")
respons_contratada = st.sidebar.text_area("Responsabilidades da Contratada (separe por ;)")
respons_contratante = st.sidebar.text_area("Responsabilidades da Contratante (separe por ;)")
texto_conclusao = st.sidebar.text_area("Texto de Conclusão")

# Upload do template
st.sidebar.subheader("📂 Upload do Template")
template_file = st.sidebar.file_uploader(
    "Carregue o modelo (.docx)",
    type=["docx"]
)

# Tabela dinâmica
st.sidebar.subheader("Tabela de Inclusões")
num_itens = st.sidebar.number_input("Quantidade de itens", min_value=0, step=1)
tabela_itens = []
for i in range(num_itens):
    item = st.sidebar.text_input(f"Item {i+1}")
    incluso = st.sidebar.text_input(f"Incluso {i+1}")
    nao_incluso = st.sidebar.text_input(f"Não Incluso {i+1}")
    tabela_itens.append({"Item": item, "Incluso": incluso, "Nao_Incluso": nao_incluso})

# Data completa
data_completa = datetime.datetime.now().strftime("%d/%m/%Y")

dados_proposta = {
    "NOME_CLIENTE": nome_cliente,
    "TITULO_PROJETO": titulo_projeto,
    "VALOR_TOTAL": valor_total,
    "PRAZO_ENTREGA": prazo_entrega,
    "ESCOPO_TECNICO": escopo_tecnico,
    "ANO": ano,
    "OBJETIVO": objetivo,
    "BENEFICIOS": beneficios,
    "REFERÊNCIAS": referencias,
    "ESCOPO": escopo,
    "OBSERVACOES": observacoes,
    "RESPONSABILIDADES_CONTRATADA": respons_contratada,
    "RESPONSABILIDADES_CONTRATANTE": respons_contratante,
    "TEXTO_CONCLUSAO": texto_conclusao,
    "DATA_COMPLETA": data_completa
}

# ==========================================================
# GERAÇÃO
# ==========================================================
campos_obrigatorios = all([nome_cliente, titulo_projeto, valor_total, prazo_entrega, escopo_tecnico])

if campos_obrigatorios:
    if st.button("🚀 Gerar Proposta"):
        arquivo_docx = gerar_docx(dados_proposta, tabela_itens, template_file)
        st.download_button(
            label="⬇️ Baixar DOCX",
            data=arquivo_docx,
            file_name=f"Proposta_{nome_cliente}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
else:
    st.warning("Preencha os campos obrigatórios (Cliente, Projeto, Valor, Prazo, Escopo Técnico) para gerar a proposta.")
